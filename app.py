import io
import re
from datetime import date
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(page_title="APA Checker + Formatter", layout="wide")
st.title("APA Reference Checker + Formatter (DOCX)")

# -----------------------------
# Citation detection (best-effort)
# -----------------------------
PAREN_CIT_RE = re.compile(r"\((?P<authors>[^()]+?),\s*(?P<year>\d{4}|n\.d\.)\)")
NARR_CIT_RE = re.compile(r"\b(?P<author>[A-Z][A-Za-z’'\-]+)\s*\((?P<year>\d{4}|n\.d\.)\)")

def normalize_year(y: str) -> str:
    y = (y or "").strip().lower()
    if y == "n.d.":
        return "n.d."
    return y if re.fullmatch(r"\d{4}", y) else "n.d."

def normalize_author_token(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"[^A-Za-z’'\- ]", "", s).strip()
    if not s:
        return "unknown"
    return s.split()[0].lower()

def extract_intext_citations(all_text: str):
    keys = set()

    # Parenthetical
    for m in PAREN_CIT_RE.finditer(all_text):
        authors = m.group("authors")
        year = normalize_year(m.group("year"))
        first = authors.split("&")[0]
        first = first.split(" and ")[0]
        first = first.split("et al.")[0]
        keys.add((normalize_author_token(first), year))

    # Narrative
    for m in NARR_CIT_RE.finditer(all_text):
        author = m.group("author")
        year = normalize_year(m.group("year"))
        keys.add((normalize_author_token(author), year))

    return keys

# -----------------------------
# References section parsing (best-effort)
# -----------------------------
def find_references_start(paragraphs):
    for i, p in enumerate(paragraphs):
        if p.text.strip().lower() == "references":
            return i
    return None

def split_reference_entries(ref_paragraphs):
    entries = []
    current = ""

    def looks_like_new_entry(t: str) -> bool:
        t = t.strip()
        if not t:
            return False
        if re.match(r"^[A-Z][A-Za-z’'\-]+,\s", t):
            return True
        if re.match(r"^[A-Z]", t) and re.search(r"\(\s*(\d{4}|n\.d\.)\s*\)", t[:80]):
            return True
        return False

    for p in ref_paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if looks_like_new_entry(t):
            if current:
                entries.append(current.strip())
            current = t
        else:
            current = (current + " " + t).strip()

    if current:
        entries.append(current.strip())

    return entries

def extract_reference_keys(reference_entries):
    keys = set()
    detailed = []

    for entry in reference_entries:
        ym = re.search(r"\(\s*(\d{4}|n\.d\.)\s*\)", entry)
        year = normalize_year(ym.group(1) if ym else "n.d.")

        am = re.match(r"^([A-Z][A-Za-z’'\-]+),\s", entry)
        if am:
            author_tok = normalize_author_token(am.group(1))
        else:
            pre = entry.split("(")[0]
            author_tok = normalize_author_token(pre)

        k = (author_tok, year)
        keys.add(k)
        detailed.append({"entry": entry, "key": k})

    return keys, detailed

# -----------------------------
# DOCX formatting helpers
# -----------------------------
def set_document_margins(doc: Document, inches=1.0):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)

def set_default_font(doc: Document, font_name: str, font_size_pt: int):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def add_page_number_top_right(doc: Document):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Clear existing runs
        for r in list(p.runs):
            r.text = ""

        run = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def apply_body_paragraph_formatting(doc: Document, double_space=True, first_line_indent_in=0.5):
    for p in doc.paragraphs:
        # Double spacing everywhere by default
        if double_space:
            p.paragraph_format.line_spacing = 2.0

        # Indent body paragraphs but avoid headings
        style_name = (p.style.name if p.style else "") or ""
        if "Heading" not in style_name and p.text.strip():
            p.paragraph_format.first_line_indent = Inches(first_line_indent_in)

def apply_references_hanging_indent(doc: Document):
    idx = find_references_start(doc.paragraphs)
    if idx is None:
        return False

    # Center the "References" heading
    heading = doc.paragraphs[idx]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.line_spacing = 2.0
    heading.paragraph_format.first_line_indent = None
    heading.paragraph_format.left_indent = None

    # Apply hanging indent to subsequent paragraphs until end
    for p in doc.paragraphs[idx + 1 :]:
        if not p.text.strip():
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    return True

def remove_first_line_indent_in_references_heading_block(doc: Document):
    """Optional: ensure references entries don't also get body first-line indent applied."""
    idx = find_references_start(doc.paragraphs)
    if idx is None:
        return
    for p in doc.paragraphs[idx + 1 :]:
        if not p.text.strip():
            continue
        # hanging indent already sets these; just ensure no extra indent remains
        # (left indent + negative first line indent is correct)
        pass

# -----------------------------
# UI
# -----------------------------
uploaded = st.file_uploader("Upload a Word document (.docx)", type=["docx"])

mode = st.radio(
    "What do you want to do?",
    ["Check references only", "Format to APA + download (and optionally check)"],
    index=0,
)

show_debug = st.checkbox("Show debug details", value=False)

st.divider()
st.subheader("Formatting options (used in Format mode)")

font_choice = st.selectbox("Font", ["Times New Roman 12", "Calibri 11", "Arial 11"], index=0)
add_pnums = st.checkbox("Add page numbers (top-right)", value=True)
format_refs = st.checkbox("Format References section (hanging indent + double spaced)", value=True)
also_check = st.checkbox("Also run reference/citation checks", value=True)

run = st.button("Run", type="primary")

if run:
    if not uploaded:
        st.error("Please upload a .docx file first.")
        st.stop()

    doc = Document(io.BytesIO(uploaded.getvalue()))
    paragraphs = list(doc.paragraphs)
    full_text = "\n".join(p.text for p in paragraphs)

    # Always compute checks if requested (or needed)
    report = None
    if mode == "Check references only" or also_check:
        cited_keys = extract_intext_citations(full_text)
        ref_start = find_references_start(paragraphs)

        if ref_start is None:
            st.warning("No 'References' heading found. I can’t reliably check the reference list.")
            st.info("Tip: Add a heading that is exactly 'References' at the end of the document.")
            if mode == "Check references only":
                st.stop()
        else:
            ref_entries = split_reference_entries(paragraphs[ref_start + 1 :])
            ref_keys, ref_details = extract_reference_keys(ref_entries)

            cited_not_in_refs = sorted(list(cited_keys - ref_keys))
            refs_not_cited = sorted(list(ref_keys - cited_keys))

            report = {
                "cited_count": len(cited_keys),
                "ref_count": len(ref_entries),
                "cited_not_in_refs": cited_not_in_refs,
                "refs_not_cited": refs_not_cited,
                "ref_details": ref_details,
                "cited_keys": sorted(list(cited_keys)),
            }

    # CHECK ONLY MODE
    if mode == "Check references only":
        if report is None:
            st.info("No check report available.")
            st.stop()

        m1, m2, m3 = st.columns(3)
        m1.metric("In-text citations found", report["cited_count"])
        m2.metric("Reference entries found", report["ref_count"])
        m3.metric("Potential issues", len(report["cited_not_in_refs"]) + len(report["refs_not_cited"]))

        st.divider()
        st.subheader("Results (best-effort)")

        if not report["cited_not_in_refs"] and not report["refs_not_cited"]:
            st.success("Looks consistent: every in-text citation appears to have a matching reference, and vice versa.")
        else:
            if report["cited_not_in_refs"]:
                st.error("Cited in text but NOT found in References (possible missing reference):")
                st.write(report["cited_not_in_refs"])
                st.caption("Keys shown as (first-author-token, year). Example: ('smith', '2021')")

            if report["refs_not_cited"]:
                st.warning("In References but NOT cited in text (possible unused reference):")
                st.write(report["refs_not_cited"])

        if show_debug and report is not None:
            st.divider()
            st.subheader("Debug details")
            st.write("**Cited keys:**", report["cited_keys"][:400])
            st.write("**Reference entries (parsed):**")
            for r in report["ref_details"][:200]:
                st.markdown(f"- `{r['key']}` → {r['entry']}")
        st.stop()

    # FORMAT MODE
    # Apply APA-ish formatting
    set_document_margins(doc, 1.0)

    if "Times New Roman" in font_choice:
        set_default_font(doc, "Times New Roman", 12)
    elif "Calibri" in font_choice:
        set_default_font(doc, "Calibri", 11)
    else:
        set_default_font(doc, "Arial", 11)

    apply_body_paragraph_formatting(doc, double_space=True, first_line_indent_in=0.5)

    if add_pnums:
        add_page_number_top_right(doc)

    formatted_refs_applied = False
    if format_refs:
        formatted_refs_applied = apply_references_hanging_indent(doc)

    # Write output
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    st.success("Formatted DOCX generated.")
    st.download_button(
        "Download APA-formatted document (.docx)",
        data=out.getvalue(),
        file_name="paper_APA_formatted.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # Show check report (optional)
    if also_check and report is not None:
        st.divider()
        st.subheader("Citation / Reference Check (best-effort)")

        m1, m2, m3 = st.columns(3)
        m1.metric("In-text citations found", report["cited_count"])
        m2.metric("Reference entries found", report["ref_count"])
        m3.metric("Potential issues", len(report["cited_not_in_refs"]) + len(report["refs_not_cited"]))

        if not report["cited_not_in_refs"] and not report["refs_not_cited"]:
            st.success("Looks consistent: every in-text citation appears to have a matching reference, and vice versa.")
        else:
            if report["cited_not_in_refs"]:
                st.error("Cited in text but NOT found in References:")
                st.write(report["cited_not_in_refs"])
            if report["refs_not_cited"]:
                st.warning("In References but NOT cited in text:")
                st.write(report["refs_not_cited"])

        if format_refs and not formatted_refs_applied:
            st.info("Note: I didn’t find a 'References' heading to format. (Formatting still applied to the rest of the document.)")
